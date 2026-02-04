"""
Шаблон для создания аналитических справок.
Использует единые стандарты форматирования.

Использование:
    exec(open('/mnt/skills/user/analytical-brief/scripts/create_brief.py').read())
    
    add_title('АНАЛИТИЧЕСКАЯ СПРАВКА')
    add_subtitle('О предмете анализа')
    add_date('02 февраля 2026 г.')
    
    add_section_heading('1. Раздел')
    add_para('Текст.')
    
    doc.save('/mnt/user-data/outputs/document.docx')
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === КОНСТАНТЫ ===
HEADER_COLOR = RGBColor(35, 50, 100)  # #233264
TABLE_HEADER_BG = "233264"
TABLE_HEADER_TEXT = RGBColor(255, 255, 255)

# === ИНИЦИАЛИЗАЦИЯ ДОКУМЕНТА ===
doc = Document()

# Настройка страницы A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Настройка базового стиля
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.15
style_normal.paragraph_format.space_after = Pt(8)
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# === ФУНКЦИИ ФОРМАТИРОВАНИЯ ===

def add_title(text, size=16):
    """Добавляет заголовок документа (по центру, цветной)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = HEADER_COLOR
    return p


def add_subtitle(text, size=12):
    """Добавляет подзаголовок (по центру, обычный)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_date(text):
    """Добавляет дату документа (по центру, курсив)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p


def add_section_heading(text, size=14):
    """Добавляет заголовок раздела (цветной, полужирный)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = HEADER_COLOR
    return p


def add_para(text, bold_parts=None):
    """
    Добавляет параграф текста.
    
    Для выделения части текста жирным используйте маркер |||:
        add_para('Обычный текст |||важная часть||| продолжение.', [1])
    
    bold_parts — список индексов частей для выделения (0, 1, 2, ...)
    """
    p = doc.add_paragraph()
    if bold_parts is None:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
    else:
        parts = text.split('|||')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            if i in bold_parts:
                run.bold = True
    return p


def add_bullet(text):
    """Добавляет пункт маркированного списка (чёрная точка)"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_numbered_item(text, indent_cm=1.5):
    """Добавляет нумерованный пункт с отступом (формат: 1) 2) 3))"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def set_cell_shading(cell, color):
    """Устанавливает цвет фона ячейки таблицы"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table(data):
    """
    Добавляет таблицу с форматированием.
    
    data — список списков: первая строка становится заголовком.
    Пример:
        data = [
            ['Показатель', 'Значение'],
            ['Строка 1', '100'],
            ['Строка 2', '200'],
        ]
        add_table(data)
    """
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    if i == 0:  # Шапка таблицы
                        run.bold = True
                        run.font.color.rgb = TABLE_HEADER_TEXT
            
            if i == 0:  # Заливка шапки
                set_cell_shading(cell, TABLE_HEADER_BG)
    
    return table


def add_empty_line():
    """Добавляет пустую строку"""
    doc.add_paragraph()


def add_signature_date(text):
    """Добавляет дату в конце документа (справа, курсив)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p


# === ГОТОВО К ИСПОЛЬЗОВАНИЮ ===
# Документ инициализирован, функции доступны.
# Используйте add_title(), add_section_heading(), add_para() и т.д.
# В конце вызовите: doc.save('/mnt/user-data/outputs/filename.docx')
